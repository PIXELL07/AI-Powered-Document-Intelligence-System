"""
Section 1: Document Ingestion and Format Normalisation.

Detects format from file bytes/extension, routes to the correct extractor,
and returns a normalized structure:

{
  "raw_text": str,
  "sections": [{"heading": str|None, "level": int, "text": str}],
  "tables": [{"sheet": str|None, "rows": [[...]]}],
  "is_scanned": bool,
  "ocr_confidence": float|None,
}

Structure (headings/sections/tables) is preserved rather than flattened,
per Section 1 requirement.
"""
import os
from app.pipeline.ocr import run_ocr_on_pdf, run_ocr_on_image


def detect_format(filename: str) -> str:
    ext = os.path.splitext(filename)[1].lower().lstrip(".")
    if ext in ("jpg", "jpeg", "png"):
        return "image"
    if ext in ("pdf",):
        return "pdf"
    if ext in ("docx",):
        return "docx"
    if ext in ("xlsx", "xls"):
        return "xlsx"
    raise ValueError(f"Unsupported file format: .{ext}")


def normalize_document(filepath: str, filename: str) -> dict:
    fmt = detect_format(filename)
    if fmt == "pdf":
        return _normalize_pdf(filepath)
    if fmt == "docx":
        return _normalize_docx(filepath)
    if fmt == "xlsx":
        return _normalize_xlsx(filepath)
    if fmt == "image":
        return _normalize_image(filepath)
    raise ValueError(f"No normalizer for format {fmt}")


def _pdf_has_text_layer(filepath: str) -> bool:
    import fitz  # PyMuPDF
    doc = fitz.open(filepath)
    try:
        for page in doc:
            if len(page.get_text("text").strip()) > 20:
                return True
        return False
    finally:
        doc.close()


def _normalize_pdf(filepath: str) -> dict:
    if _pdf_has_text_layer(filepath):
        return _extract_digital_pdf(filepath)
    # Scanned PDF: rasterize pages and OCR them.
    ocr_result = run_ocr_on_pdf(filepath)
    return {
        "raw_text": ocr_result["text"],
        "sections": _sectionize_plain_text(ocr_result["text"]),
        "tables": [],
        "is_scanned": True,
        "ocr_confidence": ocr_result["confidence"],
    }


def _extract_digital_pdf(filepath: str) -> dict:
    import fitz

    doc = fitz.open(filepath)
    sections = []
    tables = []
    full_text_parts = []
    try:
        for page_num, page in enumerate(doc):
            blocks = page.get_text("dict")["blocks"]
            for block in blocks:
                if block.get("type") != 0:
                    continue  # image block, skip for text structure
                block_text_lines = []
                max_size = 0.0
                for line in block.get("lines", []):
                    line_text = "".join(span["text"] for span in line["spans"])
                    for span in line["spans"]:
                        max_size = max(max_size, span.get("size", 0))
                    if line_text.strip():
                        block_text_lines.append(line_text)
                block_text = "\n".join(block_text_lines).strip()
                if not block_text:
                    continue
                full_text_parts.append(block_text)
                # Heuristic: larger/short single-line text = heading
                is_heading = max_size >= 12.5 and len(block_text) < 90 and "\n" not in block_text
                sections.append({
                    "heading": block_text if is_heading else None,
                    "level": 1 if is_heading else 0,
                    "text": "" if is_heading else block_text,
                    "page": page_num + 1,
                })
            # Simple table detection via PyMuPDF's table finder
            try:
                found = page.find_tables()
                for t in found.tables:
                    tables.append({"sheet": f"page_{page_num + 1}", "rows": t.extract()})
            except Exception:
                pass
    finally:
        doc.close()

    return {
        "raw_text": "\n".join(full_text_parts),
        "sections": sections,
        "tables": tables,
        "is_scanned": False,
        "ocr_confidence": None,
    }


def _normalize_docx(filepath: str) -> dict:
    import docx

    d = docx.Document(filepath)
    sections = []
    full_text_parts = []
    for para in d.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        style = (para.style.name or "").lower()
        is_heading = style.startswith("heading") or style == "title"
        level = 0
        if is_heading:
            digits = "".join(c for c in style if c.isdigit())
            level = int(digits) if digits else 1
        sections.append({
            "heading": text if is_heading else None,
            "level": level,
            "text": "" if is_heading else text,
        })
        full_text_parts.append(text)

    tables = []
    for i, table in enumerate(d.tables):
        rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
        tables.append({"sheet": f"table_{i + 1}", "rows": rows})

    return {
        "raw_text": "\n".join(full_text_parts),
        "sections": sections,
        "tables": tables,
        "is_scanned": False,
        "ocr_confidence": None,
    }


HEADER_KEYWORDS = {"desc", "description", "item", "qty", "quantity", "price", "rate",
                    "unit", "amount", "total", "date", "sku", "code"}


def _find_header_row_index(rows: list[list[str]]) -> int | None:
    """A sheet's real column-header row often isn't row 0 -- invoices
    typically have a metadata preamble ('Invoice Number:', 'Vendor:')
    above the actual line-item table. Find the first row containing at
    least 2 cells matching common tabular header vocabulary."""
    for i, row in enumerate(rows):
        matches = sum(1 for cell in row if any(kw in cell.lower() for kw in HEADER_KEYWORDS))
        if matches >= 2:
            return i
    return None


def _join_row_for_text(row: list[str]) -> str:
    non_empty = [c for c in row if c]
    # Two-cell rows are almost always "Label: Value" metadata (Invoice
    # Number, Vendor, Due Date...) -- join with a single space so
    # extraction regexes that expect "Label: Value" adjacency (no " | "
    # in between) still match. Wider rows are genuine tabular data, kept
    # pipe-delimited for readability.
    if len(non_empty) == 2:
        return " ".join(non_empty)
    return " | ".join(row)


def _normalize_xlsx(filepath: str) -> dict:
    """Identifies which sheets look like tabular financial/contractual data
    vs metadata/chart sheets, per Section 1's Excel requirement."""
    import openpyxl

    wb = openpyxl.load_workbook(filepath, data_only=True)
    tables = []
    full_text_parts = []
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        rows = []
        numeric_cells = 0
        total_cells = 0
        for row in ws.iter_rows(values_only=True):
            if all(c is None for c in row):
                continue
            rows.append(["" if c is None else str(c) for c in row])
            for c in row:
                if c is not None:
                    total_cells += 1
                    if isinstance(c, (int, float)):
                        numeric_cells += 1
        if not rows:
            continue
        # A sheet with a healthy proportion of numeric cells and multiple
        # rows/columns is treated as tabular data; sparse/text-only sheets
        # (cover pages, notes, chart-only sheets) are treated as metadata
        # and excluded from structured extraction but kept in raw text.
        numeric_ratio = (numeric_cells / total_cells) if total_cells else 0
        is_data_sheet = len(rows) >= 2 and len(rows[0]) >= 2 and numeric_ratio >= 0.15
        if is_data_sheet:
            header_idx = _find_header_row_index(rows)
            table_rows = rows[header_idx:] if header_idx is not None else rows
            tables.append({"sheet": sheet_name, "rows": table_rows})
        full_text_parts.append(f"[Sheet: {sheet_name}]\n" + "\n".join(_join_row_for_text(r) for r in rows))

    return {
        "raw_text": "\n\n".join(full_text_parts),
        "sections": [],
        "tables": tables,
        "is_scanned": False,
        "ocr_confidence": None,
    }


def _normalize_image(filepath: str) -> dict:
    ocr_result = run_ocr_on_image(filepath)
    return {
        "raw_text": ocr_result["text"],
        "sections": _sectionize_plain_text(ocr_result["text"]),
        "tables": [],
        "is_scanned": True,
        "ocr_confidence": ocr_result["confidence"],
    }


def _sectionize_plain_text(text: str) -> list[dict]:
    """OCR output has no style info to detect headings from, so fall back
    to a simple heuristic: short, all-caps or title-case standalone lines
    followed by longer text are treated as headings."""
    sections = []
    lines = [l.strip() for l in text.split("\n") if l.strip()]
    buffer = []
    for line in lines:
        looks_like_heading = len(line) < 70 and (line.isupper() or line.istitle()) and len(line.split()) <= 8
        if looks_like_heading:
            if buffer:
                sections.append({"heading": None, "level": 0, "text": " ".join(buffer)})
                buffer = []
            sections.append({"heading": line, "level": 1, "text": ""})
        else:
            buffer.append(line)
    if buffer:
        sections.append({"heading": None, "level": 0, "text": " ".join(buffer)})
    return sections
